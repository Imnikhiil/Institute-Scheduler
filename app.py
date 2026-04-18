from flask import Flask, render_template, request, send_file, jsonify
import json
import os
import traceback
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_ORIENT
from docx.enum.table import WD_ROW_HEIGHT_RULE, WD_CELL_VERTICAL_ALIGNMENT

app = Flask(__name__)
JSON_FILE = 'data.json'
LOGO_PATH = 'dseu logo.png' 

def load_config():
    if os.path.exists(JSON_FILE):
        with open(JSON_FILE, 'r') as f:
            return json.load(f)
    return {"MASTER_COURSE_DATA": {}, "TIME_SLOTS": [], "DAYS": []}

@app.route('/')
def index():
    return render_template('index.html', config=load_config())

@app.route('/generate', methods=['POST'])
def generate():
    try:
        ui_data = request.json
        entries = ui_data.get('entries', [])
        config = load_config()
        
        MASTER_DATA = config["MASTER_COURSE_DATA"]
        SLOTS = config["TIME_SLOTS"]
        DAYS = config["DAYS"]

        doc = Document()
        section = doc.sections[0]
        section.orientation = WD_ORIENT.LANDSCAPE
        section.page_width, section.page_height = section.page_height, section.page_width
        section.top_margin = section.bottom_margin = section.left_margin = section.right_margin = Inches(0.3)

        # --- HEADER SECTION (3-COLUMN TABLE FOR PERFECT CENTERING) ---
        # Column 1: Logo | Column 2: Titles | Column 3: Spacer
        header_table = doc.add_table(rows=1, cols=3)
        header_table.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Width settings (Total ~10 inches for Landscape)
        header_table.columns[0].width = Inches(1.5)
        header_table.columns[1].width = Inches(7.0)
        header_table.columns[2].width = Inches(1.5)

        # 1. Left Cell: Logo
        if os.path.exists(LOGO_PATH):
            try:
                left_cell = header_table.rows[0].cells[0]
                para_logo = left_cell.paragraphs[0]
                run_logo = para_logo.add_run()
                run_logo.add_picture(LOGO_PATH, height=Inches(0.85))
                para_logo.alignment = WD_ALIGN_PARAGRAPH.LEFT
            except Exception as img_err:
                print(f"--- IMAGE ERROR: {img_err} ---")

        # 2. Middle Cell: Titles (Perfect Center)
        mid_cell = header_table.rows[0].cells[1]
        titles = [
            "Revised Class Time Table, Odd Semester, AY-2025-26", 
            "DSEU RAJOKRI CAMPUS", 
            "Diploma in Computer Engineering, Section -B", 
            "5th Semester"
        ]
        
        for i, t in enumerate(titles):
            p = mid_cell.paragraphs[0] if i == 0 else mid_cell.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.space_after = Pt(0)
            run = p.add_run(t)
            run.bold = True
            run.font.size = Pt(10)

        # 3. Right Cell: Khali (Spacer to maintain center alignment)
        # Is cell mein kuch nahi dalna, ye balancing ke liye hai.

        doc.add_paragraph().paragraph_format.space_after = Pt(5)

        # --- TIMETABLE GRID ---
        table = doc.add_table(rows=11, cols=len(SLOTS) + 1)
        table.style = 'Table Grid'
        
        hdr = table.rows[0].cells
        hdr[0].text = "Time\nDay"
        roman = ['I','II','III','IV','V','VI','VII','VIII','IX']
        for i, slot in enumerate(SLOTS):
            cell = hdr[i+1]
            p = cell.paragraphs[0]; p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(f"{roman[i]}\n{slot}")
            run.bold = True; run.font.size = Pt(7.5)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER

        day_map = {"Monday": (1,2), "Tuesday": (3,4), "Wednesday": (5,6), "Thursday": (7,8), "Friday": (9,10)}
        for day, (rs, re) in day_map.items():
            if day in DAYS:
                cell = table.cell(rs, 0).merge(table.cell(re, 0))
                run = cell.paragraphs[0].add_run(day); run.bold = True
                cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER

        used_subs = set()
        for e in entries:
            used_subs.add(e['sub_full'])
            rs, re = day_map[e['day']]
            col_idx = SLOTS.index(e['start_t']) + 1
            duration = SLOTS.index(e['end_t']) - SLOTS.index(e['start_t']) + 1
            
            is_lec = e['is_lec']
            t_row = rs if (is_lec or e['pos'] == "Upper") else re
            target = table.cell(t_row, col_idx)
            
            if is_lec: target = target.merge(table.cell(re, col_idx))
            if duration > 1:
                for d in range(1, duration):
                    side = table.cell(t_row, col_idx + d)
                    if is_lec: side = side.merge(table.cell(re, col_idx + d))
                    target = target.merge(side)

            target.text = f"{e['acr']} {e['room']}"
            target.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            target.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER

        for idx, row in enumerate(table.rows):
            row.height_rule = WD_ROW_HEIGHT_RULE.EXACTLY
            row.height = Pt(18) if idx == 0 else Pt(14.40)
            for cell in row.cells:
                for p in cell.paragraphs:
                    p.paragraph_format.line_spacing = 1.0
                    for run in p.runs: run.font.size = Pt(8)

        # Footer Table
        doc.add_paragraph().add_run("\n")
        ft = doc.add_table(rows=1, cols=5); ft.style = 'Table Grid'
        cols = ["Semester", "Course Name", "Credit", "Accronym", "Faculty"]
        for i, txt in enumerate(cols):
            ft.rows[0].cells[i].text = txt
            ft.rows[0].cells[i].paragraphs[0].runs[0].bold = True

        for s in sorted(used_subs):
            info = MASTER_DATA[s]
            r = ft.add_row().cells
            for i, val in enumerate([info[0], s, info[1], info[2], info[3]]):
                r[i].text = str(val)
                r[i].paragraphs[0].runs[0].font.size = Pt(8.5)

        output_file = "Timetable_Generated.docx"
        doc.save(output_file)
        return send_file(output_file, as_attachment=True)

    except Exception as e:
        print("--- SERVER CRASHED ---")
        traceback.print_exc() 
        return jsonify({"error": str(e)}), 500

if __name__ == '__main__':
    app.run(debug=True)